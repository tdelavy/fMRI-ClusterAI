import streamlit as st
import streamlit.components.v1 as components
import subprocess
import tempfile
import os
from openai import OpenAI
import re
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import nibabel as nib
import numpy as np
import matplotlib.pyplot as plt
import io
from matplotlib.colors import ListedColormap, BoundaryNorm
from skimage.measure import marching_cubes
import plotly.graph_objects as go
import json, urllib.parse
import pandas as pd

st.set_page_config(page_title="AItlas Clusters", page_icon="🧠")

welcome_paragraph = """
<h4 style="color:#0c5460;margin-bottom:5px;">🎉 Welcome back!</h4>
<p style="color:#0c5460;margin:0;">
    We've added <b>new Neurosynth visualizations</b> for each cluster, allowing you to explore the functional connectivity of your specific clusters at rest. Brodmann atlas and a 3D visualization tool are also integrated in Pro Mode!
</p>
"""

# Instantiate the client with the Perplexity API endpoint.
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    st.error("Please set your OPENAI_API_KEY environment variable!")
else:
    client = OpenAI(api_key=openai_api_key, base_url="https://api.perplexity.ai")

def load_label_file(file_path):
    """
    Reads a label file (in the format produced by AFNI's 3dinfo -labeltable)
    and returns a dictionary mapping integer labels to region names.
    
    Expected file format (each line):
      "9" "right_Area_hOc4v_(LingG)"
      "10" "right_Area_hOc3v_(LingG)"
      ...
    """
    label_dict = {}
    with open(file_path, "r") as f:
        for line in f:
            line = line.strip()
            if not line.startswith('"'):
                continue
            matches = re.findall(r'"(.*?)"', line)
            if len(matches) >= 2:
                try:
                    key = int(matches[0])
                    label_dict[key] = matches[1]
                except ValueError:
                    continue
    return label_dict

# Load label dictionaries from your text files.
label_julich = load_label_file("labelJulich.txt")
label_fs_afni = load_label_file("labelFs.afni.txt")
label_Brodmann = load_label_file("labelBrodmann.txt")

def search_all_labels(voxel, atlas_data, max_radius=7):
    """
    Searches around the given voxel (within a sphere of radius max_radius)
    for all nonzero labels and returns a dictionary mapping label values
    to their minimum Euclidean distance (in voxel units) from the original voxel.
    """
    shape = atlas_data.shape
    label_distances = {}
    for di in range(-max_radius, max_radius+1):
        for dj in range(-max_radius, max_radius+1):
            for dk in range(-max_radius, max_radius+1):
                # Only consider offsets within a sphere of radius max_radius.
                if di**2 + dj**2 + dk**2 <= max_radius**2:
                    ni = voxel[0] + di
                    nj = voxel[1] + dj
                    nk = voxel[2] + dk
                    if 0 <= ni < shape[0] and 0 <= nj < shape[1] and 0 <= nk < shape[2]:
                        lab = int(atlas_data[ni, nj, nk])
                        if lab != 0:
                            dist = (di**2 + dj**2 + dk**2)**0.5
                            if lab in label_distances:
                                label_distances[lab] = min(label_distances[lab], dist)
                            else:
                                label_distances[lab] = dist
    return label_distances

def get_anatomical_labels(coord, atlas, coord_system="RAI"):
    """
    Maps a coordinate (x, y, z) in MNI space to the correct voxel index using NiBabel's affine,
    then returns the label name plus the voxel index. If no label is found, returns an appropriate message.
    """

    # 1) Determine atlas path & label dictionary
    if atlas == "Julich_MNI2009c":
        atlas_path = "Julich_MNI2009c.nii.gz"
        label_dict = label_julich
    elif atlas == "FreeSurfer_MNI2009c":
        atlas_path = "FS.afni.MNI2009c_asym.nii.gz"
        label_dict = label_fs_afni
    elif atlas == "Broadmann_pijn":
        atlas_path = "Brodmann_pijn_afni.nii.gz"
        label_dict = label_Brodmann
    else:
        return "Unknown atlas"

    # 2) Load the atlas
    try:
        atlas_img = nib.load(atlas_path)
    except Exception as e:
        return f"Error loading atlas: {e}"
    atlas_data = atlas_img.get_fdata()
    shape = atlas_data.shape

    # 3) Convert coordinate system if needed
    if coord_system.upper().startswith("LPI"):
        # Coordinates are already in LPI; use them directly.
        x_mni, y_mni, z_mni = coord
    elif coord_system.upper() == "RAI":
        # Convert RAI to LPI by inverting x and y.
        x_mni, y_mni, z_mni = -coord[0], -coord[1], coord[2]
    else:
        # If unknown, assume coordinates are already in MNI.
        x_mni, y_mni, z_mni = coord
        
    neurosynth_url = f"https://neurosynth.org/locations/?x={int(round(x_mni))}&y={int(round(y_mni))}&z={int(round(z_mni))}"

    # 4) Use the atlas's affine to get voxel indices
    affine = atlas_img.affine
    inv_affine = np.linalg.inv(affine)
    voxel_coord = inv_affine.dot([x_mni, y_mni, z_mni, 1])[:3]
    i, j, k = np.round(voxel_coord).astype(int)

    # 5) Check bounds
    if not (0 <= i < shape[0] and 0 <= j < shape[1] and 0 <= k < shape[2]):
        return "Coordinate outside atlas volume"

    # 6) Direct label lookup
    label_val = int(atlas_data[i, j, k])
    if label_val != 0:
        region_name = label_dict.get(label_val, f"Label {label_val} not found")
        return region_name, (i, j, k), neurosynth_url, (x_mni, y_mni, z_mni)
    else:
        # 7) If 0, search nearby within 7 mm
        label_candidates = search_all_labels((i, j, k), atlas_data, max_radius=7)
        if label_candidates:
            sorted_candidates = sorted(label_candidates.items(), key=lambda x: x[1])
            search_info = "\n".join([
                f"* Within {round(dist,1)} mm: {label_dict.get(lab, 'Label '+str(lab))}"
                for lab, dist in sorted_candidates
            ])
            return search_info, (i, j, k), neurosynth_url, (x_mni, y_mni, z_mni)
        else:
            return "Label 0 not found"

def parse_cluster_file(file_path, max_clusters=6):
    clusters = []
    with open(file_path, "r") as f:
        for line in f:
            line = line.strip()
            if line.startswith("#") or not line:
                continue
            parts = line.split()
            if len(parts) >= 7:
                # AFNI format: [vox, CM_x, CM_y, CM_z, Peak_x, Peak_y, Peak_z, ...]
                cluster = {
                    'voxels': int(parts[0]),
                    'CM': (float(parts[1]), float(parts[2]), float(parts[3])),
                    'Peak': (float(parts[4]), float(parts[5]), float(parts[6]))
                }
            elif len(parts) == 4:
                # SPM conversion format: [vox, Peak_x, Peak_y, Peak_z]
                cluster = {
                    'voxels': int(parts[0]),
                    'CM': None,
                    'Peak': (float(parts[1]), float(parts[2]), float(parts[3]))
                }
            else:
                continue
            clusters.append(cluster)
            if len(clusters) == max_clusters:  # Only keep the first 6 clusters.
                break
    return clusters
    
def synthesize_interpretation(anatomical_info, task_description, contrast_description, max_clusters):
    """
    Sends a prompt to the Perplexity API, streams tokens, and progressively
    updates Streamlit to show partial output as it arrives.
    """
    prompt = (
        "You are a specialized neuroscience literature analysis assistant."
        f"Based on the following anatomical information for up to {max_clusters} clusters from a fMRI study.\n"
        "Please provide a detailed literature analysis as follows:\n\n"
        "1) Review each cluster information.\n"
        "2) Explain the cluster’s involvement in the fMRI task (if any relevant evidence or prior studies exist). If there is 0 known involvement, say so.\n"
        "3) Provide a detailed summary of up-to-date literature from reputable sources, if available, on that region’s role in the tasks or processes. Make sure you give enough relevant information to understand the role of the cluster in known litterature. But, if no relevant studies are found, say that evidence is currently limited or inconclusive.\n"
        "\n"
        f"Theme/domain of the research: {task_description}\n"
        f"Explanation of the clusters found: {contrast_description}\n"
        "Below is the full anatomical labels output for the relevant clusters:\n"
        f"{anatomical_info}\n\n"
        "Now synthesize a clear, accurate interpretation for each clusters."
    )
    
    # Make a SINGLE call with model="sonar-pro" or "sonar"
    response = client.chat.completions.create(
        model="sonar-deep-research", #-deep-research
        messages=[
            {"role": "user", "content": prompt}
        ],
    )

        # Extract the raw text
    interpretation = response.choices[0].message.content
    
    # Strip out any <think>...</think> segments, if present
    interpretation = re.sub(r"<think>.*?</think>", "", interpretation, flags=re.DOTALL)

    references = getattr(response, "citations", [])

    # Then extract the text
    return interpretation, references

def run_analysis(cluster_file_path, atlas, task_description, contrast_description):
    """
    Runs the full analysis:
      - Parses the cluster file (keeping only the first 6 clusters).
      - Obtains anatomical labels for each cluster.
      - Combines the outputs into a string.
      - Calls ChatGPT for interpretation.
    Returns both the raw anatomical outputs and ChatGPT's interpretation.
    """
    clusters = parse_cluster_file(cluster_file_path, max_clusters)
    anatomical_results = []
    
    # Display the first 6 clusters.
    st.subheader(f"First {max_clusters} clusters:")

    if atlas == "Julich_MNI2009c":
        atlas_path = "Julich_MNI2009c.nii.gz"
    elif atlas == "FreeSurfer_MNI2009c":
        atlas_path = "FS.afni.MNI2009c_asym.nii.gz"
    elif atlas == "Broadmann_MNI2009":
        atlas_path = "Brodmann_pijn_afni.nii.gz"

    atlas_img = nib.load(atlas_path)  # the same atlas you used in get_anatomical_labels
    atlas_data = atlas_img.get_fdata()

    # Get anatomical labels for each cluster.
    for i, cluster in enumerate(clusters, start=1):
        try:
            label_info, (vi, vj, vk), neurosynth_url, lpi_coords = get_anatomical_labels(cluster["Peak"], atlas, coord_system)
        except ValueError:
            st.warning(
                f"Cluster {i}: The analysis did not return any relevant brain location. We suspect the coordinate system your indicated is incorrect. "
                "Please verify whether your coordinates are in RAI or LPI. If you used the 1D File Creator, ensure you are using the LPI system."
            )
            continue

        # Check for common errors indicating a coordinate system issue.
        if isinstance(label_info, str) and ("Coordinate outside atlas volume" in label_info or "Label 0 not found" in label_info):
            st.warning(
                f"Cluster {i}: The analysis did not return any relevant brain location. We suspect the coordinate system your indicated is incorrect. "
                "Please verify whether your coordinates are in RAI or LPI. If you used the 1D File Creator, ensure you are using the LPI system."
            )
            continue
            
        st.write(f"**Cluster {i}:** Voxels: {cluster['voxels']}, Peak (LPI): {tuple(round(c, 2) for c in lpi_coords)}")
        st.text(label_info)

        with st.expander("🧠 Functional Resting-State Connectivity (via Neurosynth.org)"):
            components.html(f"""
                <div style="transform: scale(0.55); transform-origin: top left; width: 1000px; height: 750px;">
                    <iframe src="{neurosynth_url}" width="1600" height="950" style="border:none;"></iframe>
                </div>
            """, height=550)

        st.write("-" * 20)
        anatomical_results.append(f"Cluster {i} (Voxels: {cluster['voxels']}, Peak: {cluster['Peak']}):\n{label_info}")
        
    anatomical_str = "\n\n".join(anatomical_results)
        
    return anatomical_str, clusters

def sanitize_filename(s):
    # Replace non-word characters with underscores
    return re.sub(r'\W+', '_', s)

def create_word_report(settings_summary, anatomical_str, interpretation, references):
    """
    Creates a Word document containing the settings summary,
    the parsed clusters, and the ChatGPT interpretation.
    Returns the path to the temporary .docx file.
    """
    doc = docx.Document()

    # Title
    title = doc.add_heading('fMRI Cluster Analysis Report', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Analysis Settings Section
    doc.add_heading('Analysis Settings', level=1)
    p = doc.add_paragraph(settings_summary)
    p.style.font.size = Pt(11)

    # Cluster Summary Section (as a table)
    doc.add_heading('Cluster Summary', level=1)
    # Create a table with headers: Cluster, Voxels, Peak (LPI), Anatomical Label
    table = doc.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Cluster'
    hdr_cells[1].text = 'Voxels'
    hdr_cells[2].text = 'Peak (LPI)'
    hdr_cells[3].text = 'Anatomical Label'
    # Assume anatomical_str has each cluster on a new line in the format:
    # "Cluster i (Voxels: xxx, Peak: (x, y, z)):\nLabel info"
    for line in anatomical_str.split('\n\n'):
        # Basic parsing; you might want to improve this based on your output format
        if not line.strip():
            continue
        parts = line.split("):")
        if len(parts) < 2:
            continue
        header_line = parts[0] + ")"
        label_text = parts[1].strip()
        # Extract cluster number, voxels, and peak coordinates from header_line
        # Example header_line: "Cluster 1 (Voxels: 654, Peak: (24.5, 80.5, 29.5))"
        try:
            cluster_number = header_line.split(" ")[1]
            voxels = header_line.split("Voxels:")[1].split(",")[0].strip()
            peak = header_line.split("Peak:")[1].strip().strip("()")
        except Exception:
            cluster_number, voxels, peak = "", "", ""
        row_cells = table.add_row().cells
        row_cells[0].text = cluster_number
        row_cells[1].text = voxels
        row_cells[2].text = peak
        row_cells[3].text = label_text

    # Interpretation Section
    doc.add_heading('Interpretation', level=1)
    # Process the interpretation text to handle markdown-style formatting.
    for para in interpretation.split("\n"):
        para = para.strip()
        if not para:
            continue

        # Headings
        if para.startswith("### "):
            doc.add_heading(para[4:].strip(), level=2)
        elif para.startswith("#### "):
            doc.add_heading(para[5:].strip(), level=3)
        else:
            # Bold text within the paragraph
            p = doc.add_paragraph()
            bold_parts = re.split(r'(\*\*.*?\*\*)', para)
            for part in bold_parts:
                run = p.add_run(part.replace("**", "")) if part.startswith("**") else p.add_run(part)
                run.bold = part.startswith("**")

    # References Section
    doc.add_heading('References', level=1)
    if references:
        for i, ref in enumerate(references, start=1):
            p = doc.add_paragraph(style='List Number')
            p.add_run(f"{i}. ").bold = True
            p.add_run(ref)
    else:
        doc.add_paragraph("No references returned.", style='Normal')

    # Save to a temporary file
    temp_doc = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_doc.name)
    return temp_doc.name

# -------------------------------
# Streamlit App Layout
# -------------------------------
#st.title("fMRI Cluster Analysis with Atlas Labeling and Perplexity AI")

# Initialize Pro Mode flag if not already set
if "pro_mode" not in st.session_state:
    st.session_state.pro_mode = False

# Create the Pro Mode indicator HTML
if st.session_state.pro_mode:
    pro_status_html = (
        "<a href='?page=select' title='Pro Mode is enabled.'>"
        "<div style='display:inline-block; border-radius:50%; width:13px; height:13px; background-color: green; margin-right: 5px;'></div>"
        "</a>"
    )
else:
    pro_status_html = (
        "<a href='?page=select' title='Pro Mode is disabled.'>"
        "<div style='display:inline-block; border-radius:50%; width:13px; height:13px; background-color: red; margin-right: 5px;'></div>"
        "</a>"
    )

# Combine the header text with the indicator HTML
header_html = f"<h1>Analysis Settings {pro_status_html}</h1>"
st.sidebar.markdown(header_html, unsafe_allow_html=True)

# Define the options
if st.session_state.pro_mode:
    options = ["-- Select an option --", "1D File Creator", "Analysis"]
else:
    options = ["-- Select an option --", "1D File Creator (Pro Mode required)", "Analysis"]

# Create the selectbox
conversion_choice = st.sidebar.selectbox(
    label="Analysis",
    options=options,
    index=0,
    help="Select 'Analysis' if you already have a .1D cluster file; select '1D File Creator' if you need a .1D file."
)

# If non-Pro and Other is selected, reset and show warning.
if not st.session_state.pro_mode and conversion_choice.startswith("1D File Creator"):
    st.sidebar.warning("1D File Creator is locked. Enable Pro Mode to use this option.")
    # Reset the selection to default:
    conversion_choice = "-- Select an option --"

# Check if a valid option has been selected
if conversion_choice == "-- Select an option --":
    if not st.session_state.pro_mode:
        pro_password = st.sidebar.text_input("Enter Pro Mode Password", type="password")
        if pro_password:
            if pro_password == st.secrets["api_password"]:
                st.session_state.pro_mode = True
                st.rerun()  # Immediately update the UI
            else:
                st.sidebar.error("Incorrect Pro Mode password.")
    if st.session_state.pro_mode:
        st.sidebar.success("Pro Mode enabled!")

    if "welcome_shown" not in st.session_state:
        with st.container():
            st.markdown(f"""
<div style="background-color:#d1ecf1;padding:10px;border-radius:5px;border-left:5px solid #0c5460;">
{welcome_paragraph}
</div>
""", unsafe_allow_html=True)
        st.session_state.welcome_shown = True
                
    st.title("AItlas")
    st.subheader("Analyze your fMRI clusters with anatomical labeling and Deep Research AI from Perplexity")
    st.markdown(""" 
    #### AItlas: fMRI Cluster Analysis

    AItlas is designed to analyze fMRI clusters. It reads a `.1D` file containing the voxel numbers along with the X, Y, and Z coordinates of the peak of each cluster and suggests two atlases based on precision levels for identifying brain regions.

    The extracted anatomical information is then analyzed by Perplexity’s Sonar Deep Research, which returns relevant literature on the number of clusters selected for the specified task and condition.
    """)

    st.markdown(""" 
    <u>**Note:**</u> To use the AItlas Pro Mode, request the password
    [**here**](mailto:thibaud.delavy@bluewin.ch?subject=Request%20for%20AI%20Pro%20Mode&body=Hello%20Thibaud,%0A%0AI%20would%20like%20to%20request%20access%20to%20the%20AI%20Pro%20Mode.%0A%0ABest,%0A%5BYour%20Name%5D).
    """, unsafe_allow_html=True)

    st.markdown(""" 
    ##### 📌 **Coordinate Systems Used**
    - **RAI (used by AFNI)**: Right-To-Left, Anterior-To-Posterior, Inferior-To-Superior.
    - **LPI (used by SPM)**: Left-To-Right, Posterior-To-Anterior, Inferior-To-Superior.
    """)

    st.markdown("""
    <u>**Important Template Information:**</u>

    For optimal performance, this application assumes that your fMRI data has been normalized using the **MNI152 2009 template**. The voxel-to-anatomical label mapping is calibrated specifically for this template. If a different template was used during preprocessing, the anatomical labeling may not be accurate. Please ensure that your analysis employed the MNI152 2009 template to benefit from this app.
    """, unsafe_allow_html=True)
    st.warning("Please select either 1D File Creator or Analysis to continue.")
    st.markdown(
        """
        <style>
        .credits-discrete {
            font-size: 0.85rem; /* Slightly smaller than normal text */
            color: #888;        /* Gray text */
            text-align: right;  /* Align to the right side of its container */
            margin-right: 1rem; /* Optional spacing from the edge */
        }
        .highlight-name {
            text-decoration: underline; /* Underline the name */
        }
        </style>
        <div class="credits-discrete">
        Credits to 
        <a href="https://www.linkedin.com/in/thibaud-delavy-0b13571b9" 
            target="_blank"
            style="color: inherit; text-decoration: none;">
            <span class="highlight-name">Thibaud Delavy</span>
        </a>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.stop()

cluster_file_path = None  # This will store the path to the .1D file for analysis

if conversion_choice == "1D File Creator":
    st.title("AItlas - Custom Cluster Creation")
    st.subheader("Enter your custom peak activation coordinates")

    st.markdown(
        """
        For the analysis, the system requires a .1D file containing the cluster information.
        This page allows you to create one by entering the number of voxels and the peak activation 
        coordinates (X, Y, Z) for each cluster in the table below. Once you have filled in the required 
        information, click the download button to obtain your .1D file and proceed with the analysis.
        """
    )

    input_coord_system = st.selectbox(
         "Input Coordinate System",
         options=["LPI", "RAI"],
         index=0,
         help="Select the coordinate system for your input coordinates. Coordinates will be converted to LPI for analysis."
    )

    # Provide a data editor table with columns: Voxels, Peak x, Peak y, Peak z.
    # Here we create 12 empty rows as default so the user sees a fixed-size table.
    default_data = pd.DataFrame({
        "Voxels": [None] * 12,
        "Peak x": [None] * 12,
        "Peak y": [None] * 12,
        "Peak z": [None] * 12
    })
    st.markdown("**Please enter your custom peak activation coordinates below.**\n\n*Note: The default table contains 12 rows. Fill in only the rows you need. Rows with missing values will be skipped.*")
    custom_clusters = st.data_editor(
        default_data,
        num_rows="dynamic",
        use_container_width=True,
        key="custom_clusters"
    )

    # Function to convert coordinates to LPI if needed.
    def convert_to_lpi(row):
        # If input is in RAI, convert to LPI by negating x and y.
        if input_coord_system == "RAI":
            return pd.Series({
                "Voxels": row["Voxels"],
                "Peak x": -row["Peak x"] if pd.notnull(row["Peak x"]) else None,
                "Peak y": -row["Peak y"] if pd.notnull(row["Peak y"]) else None,
                "Peak z": row["Peak z"]
            })
        else:
            return row  # Already in LPI

    # Convert columns to numeric (turns empty or invalid entries into NaN).
    for col in ["Voxels", "Peak x", "Peak y", "Peak z"]:
        custom_clusters[col] = pd.to_numeric(custom_clusters[col], errors="coerce")

    # Apply coordinate system conversion.
    custom_clusters_lpi = custom_clusters.apply(convert_to_lpi, axis=1)

    # Create the .1D file content.
    header = (
        "#Coordinate order = LPI\n"
        "#Voxels   Peak x Peak y Peak z\n"
        "#------ ------ ------ ------\n"
    )
    rows = []
    for _, row in custom_clusters_lpi.iterrows():
        try:
            # Attempt to convert the values. If any conversion fails, skip the row.
            vox = int(row['Voxels'])
            px = float(row['Peak x'])
            py = float(row['Peak y'])
            pz = float(row['Peak z'])
            row_str = f"{vox:7d} {px:+7.1f} {py:+7.1f} {pz:+7.1f}"
            rows.append(row_str)
        except (TypeError, ValueError):
            # If conversion fails (e.g., value is None), skip this row.
            continue

    file_content = header + "\n" + "\n".join(rows)

    st.download_button(
        label="Download Custom .1D File",
        data=file_content,
        file_name="Custom_Clusters_LPI.1D",
        mime="text/plain"
    )

    st.warning(
        "The coordinates you will enter will be automatically converted to LPI format. "
        "Please ensure that 'LPI' is selected as the Coordinate System in the analysis for accurate processing."
    )

elif conversion_choice == "Analysis":
    task_description = st.sidebar.text_input(
        "Theme / Domain",
        "Cognitive fatigue",
        help="Enter the theme or domain of your study (e.g., cognitive fatigue) to give context to Perplexity"
    )
    contrast_description = st.sidebar.text_area(
        "Analysis description (e.g., contrast)",
        "",
        help="Enter the overall objective of your analysis (e.g., Incongruent minus Congruent in stroop task)"
    )
    
    # AFNI is selected: uploader for .1D file only
    uploaded_file = st.sidebar.file_uploader("Choose your cluster file (.1D)", type=["1D"])
    if uploaded_file is not None:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".1D") as tmp:
            tmp.write(uploaded_file.read())
            cluster_file_path = tmp.name
    
    coord_system = st.sidebar.selectbox(
        label="Coordinate System",
        options=["RAI", "LPI (=SPM)"],
        index=0,
        help="RAI format is used by AFNI by default, while LPI is used by SPM. Choose LPI if used the 1D File Creator."
    )
    if st.session_state.pro_mode:
        atlas = st.sidebar.selectbox(
            label="Choose your atlas",
            options=["Julich_MNI2009c", "FreeSurfer_MNI2009c", "Broadmann_MNI2009"],
            index=1,  # default "Julich_MNI2009c"
            help="Select which atlas to use for anatomical labeling"
        )
    else:
        atlas = st.sidebar.selectbox(
            label="Choose your atlas",
            options=["Julich_MNI2009c", "FreeSurfer_MNI2009c", "Broadmann_MNI2009"],
            index=1,
            disabled=True,
            help="Additional atlas options are available in Pro mode: FreeSurfer_MNI2009c & Broadmann_pijn"
        )
    if st.session_state.pro_mode:
        max_clusters = st.sidebar.number_input(
            "Max. number of clusters to include",
            min_value=1,
            max_value=12,
            value=6,
            step=1,
            help="Select how many clusters will be used in the analysis. Maximum 12 clusters."
        )
    else:
        max_clusters = st.sidebar.number_input(
            "Max. number of clusters to include",
            value=4,
            disabled=True,
            help="Only 4 clusters are available in the non pro version."
        )
    if st.session_state.pro_mode:
        use_ai = st.sidebar.checkbox(
            "Use Sonar Deep Research", 
            value=False,
            help="Pro mode allows deep literature research via Perplexity."
        )
    else:
        use_ai = st.sidebar.checkbox(
            "Use Sonar Deep Research", 
            value=False, 
            disabled=True,
            help="This feature is available only in Pro Mode."
        )

# --- Run Analysis button (only for Analysis option) ---
if conversion_choice == "Analysis":
    st.title("AItlas")
    st.subheader("Analyze your fMRI clusters with anatomical labeling and Deep Research AI from Perplexity")
    st.markdown(""" 
    AItlas is designed to analyze fMRI clusters. Download your `.1D` file containing the voxel numbers along with the X, Y, and Z coordinates of the peak of each cluster (in **RAI** or **LPI** coordinate systems) and choose one of the three atlases.

    The extracted anatomical information is then analyzed by Perplexity’s Sonar Deep Research, which returns relevant literature on the number of clusters selected for the specified task and condition.
    """)


    if atlas == "Julich_MNI2009c":
        st.image("Julich_MNI2009c_Atlas.png", caption="Julich MNI2009c Atlas")
        with st.expander("View Julich Atlas Region List"):
            st.text("""\

    Atlas Julich_MNI2009c,      314 regions
    ----------- Begin regions for Julich_MNI2009c atlas-----------
    u:right_Area_45_(IFG):1  
    u:left_Area_45_(IFG):201
    u:right_Area_44_(IFG):2  
    u:left_Area_44_(IFG):202
    u:right_Area_Fo1_(OFC):3  
    u:left_Area_Fo1_(OFC):203
    u:right_Area_Fo2_(OFC):4  
    u:left_Area_Fo2_(OFC):204
    u:right_Area_Fo3_(OFC):5  
    u:left_Area_Fo3_(OFC):205
    u:right_Area_hOc5_(LOC):6  
    u:left_Area_hOc5_(LOC):206
    u:right_Area_hOc2_(V2,_18):7  
    u:left_Area_hOc2_(V2,_18):207
    u:right_Area_hOc1_(V1,_17,_CalcS):8  
    u:left_Area_hOc1_(V1,_17,_CalcS):208
    u:right_Area_hOc4v_(LingG):9  
    u:left_Area_hOc4v_(LingG):209
    u:right_Area_hOc3v_(LingG):10 
    u:left_Area_hOc3v_(LingG):210
    u:right_Area_TE_1.0_(HESCHL):11 
    u:left_Area_TE_1.0_(HESCHL):211
    u:right_Area_TE_2.1_(STG):12 
    u:left_Area_TE_2.1_(STG):212
    u:right_Area_TPJ_(STG/SMG):13 
    u:left_Area_TPJ_(STG/SMG):213
    u:right_Area_TE_1.2_(HESCHL):14 
    u:left_Area_TE_1.2_(HESCHL):214
    u:right_Area_TE_3_(STG):15 
    u:left_Area_TE_3_(STG):215
    u:right_Area_TE_1.1_(HESCHL):16 
    u:left_Area_TE_1.1_(HESCHL):216
    u:right_Area_TE_2.2_(STG):17 
    u:left_Area_TE_2.2_(STG):217
    u:right_Area_33_(ACC):18 
    u:left_Area_33_(ACC):218
    u:right_Area_s32_(sACC):19 
    u:left_Area_s32_(sACC):219
    u:right_Area_p32_(pACC):20 
    u:left_Area_p32_(pACC):220
    u:right_Area_Id2_(Insula):21 
    u:left_Area_Id2_(Insula):221
    u:right_Area_Id3_(Insula):22 
    u:left_Area_Id3_(Insula):222
    u:right_CA2_(Hippocampus):23 
    u:left_CA2_(Hippocampus):223
    u:right_CA3_(Hippocampus):24 
    u:left_CA3_(Hippocampus):224
    u:right_Entorhinal_Cortex:25 
    u:left_Entorhinal_Cortex:225
    u:right_DG_(Hippocampus):26 
    u:left_DG_(Hippocampus):226
    u:right_CA1_(Hippocampus):27 
    u:left_CA1_(Hippocampus):227
    u:right_HATA_(Hippocampus):28 
    u:left_HATA_(Hippocampus):228
    u:right_Area_OP4_(POperc):29 
    u:left_Area_OP4_(POperc):229
    u:right_Area_OP1_(POperc):30 
    u:left_Area_OP1_(POperc):230
    u:right_Area_OP2_(POperc):31 
    u:left_Area_OP2_(POperc):231
    u:right_Area_OP3_(POperc):32 
    u:left_Area_OP3_(POperc):232
    u:right_Area_FG2_(FusG):33 
    u:left_Area_FG2_(FusG):233
    u:right_Area_FG1_(FusG):34 
    u:left_Area_FG1_(FusG):234
    u:right_Area_PGp_(IPL):35 
    u:left_Area_PGp_(IPL):235
    u:right_Area_PFt_(IPL):36 
    u:left_Area_PFt_(IPL):236
    u:right_Area_PGa_(IPL):37 
    u:left_Area_PGa_(IPL):237
    u:right_Area_PFop_(IPL):38 
    u:left_Area_PFop_(IPL):238
    u:right_Area_PFm_(IPL):39 
    u:left_Area_PFm_(IPL):239
    u:right_Area_PFcm_(IPL):40 
    u:left_Area_PFcm_(IPL):240
    u:right_Area_Ig2_(Insula):41 
    u:left_Area_Ig2_(Insula):241
    u:right_Area_Ig1_(Insula):42 
    u:left_Area_Ig1_(Insula):242
    u:right_Area_Id1_(Insula):43 
    u:left_Area_Id1_(Insula):243
    u:right_Area_hOc4lp_(LOC):44 
    u:left_Area_hOc4lp_(LOC):244
    u:right_Area_hOc4la_(LOC):45 
    u:left_Area_hOc4la_(LOC):245
    u:right_Area_hOc4d_(Cuneus):46 
    u:left_Area_hOc4d_(Cuneus):246
    u:right_Area_hOc3d_(Cuneus):47 
    u:left_Area_hOc3d_(Cuneus):247
    u:right_Area_4p_(PreCG):48 
    u:left_Area_4p_(PreCG):248
    u:right_Area_4a_(PreCG):49 
    u:left_Area_4a_(PreCG):249
    u:right_Area_1_(PostCG):50 
    u:left_Area_1_(PostCG):250
    u:right_Area_3a_(PostCG):51 
    u:left_Area_3a_(PostCG):251
    u:right_Area_3b_(PostCG):52 
    u:left_Area_3b_(PostCG):252
    u:right_Area_hIP1_(IPS):53 
    u:left_Area_hIP1_(IPS):253
    u:right_Area_hIP2_(IPS):54 
    u:left_Area_hIP2_(IPS):254
    u:right_Area_5L_(SPL):55 
    u:left_Area_5L_(SPL):255
    u:right_Area_5M_(SPL):56 
    u:left_Area_5M_(SPL):256
    u:right_Area_7PC_(SPL):57 
    u:left_Area_7PC_(SPL):257
    u:right_Area_hIP3_(IPS):58 
    u:left_Area_hIP3_(IPS):258
    u:right_Area_7A_(SPL):59 
    u:left_Area_7A_(SPL):259
    u:right_Area_7M_(SPL):60 
    u:left_Area_7M_(SPL):260
    u:right_Area_5Ci_(SPL):61 
    u:left_Area_5Ci_(SPL):261
    u:right_Area_Id7_(Insula):62 
    u:left_Area_Id7_(Insula):262
    u:right_Area_s24_(sACC):63 
    u:left_Area_s24_(sACC):263
    u:right_Area_25_(sACC):64 
    u:left_Area_25_(sACC):264
    u:right_SF_(Amygdala):65 
    u:left_SF_(Amygdala):265
    u:right_LB_(Amygdala):66 
    u:left_LB_(Amygdala):266
    u:right_Subiculum_(Hippocampus):67 
    u:left_Subiculum_(Hippocampus):267
    u:right_Area_PF_(IPL):68 
    u:left_Area_PF_(IPL):268
    u:right_Area_7P_(SPL):69 
    u:left_Area_7P_(SPL):269
    u:right_Area_Fp2_(FPole):70 
    u:left_Area_Fp2_(FPole):270
    u:right_Area_Fp1_(FPole):71 
    u:left_Area_Fp1_(FPole):271
    u:right_Fastigial_Nucleus_(Cerebellum):72 
    u:left_Fastigial_Nucleus_(Cerebellum):272
    u:right_VTM_(Amygdala):73 
    u:left_VTM_(Amygdala):273
    u:right_Area_p24ab_(pACC):74 
    u:left_Area_p24ab_(pACC):274
    u:right_Area_p24c_(pACC):75 
    u:left_Area_p24c_(pACC):275
    u:right_MF_(Amygdala):76 
    u:left_MF_(Amygdala):276
    u:right_IF_(Amygdala):77 
    u:left_IF_(Amygdala):277
    u:right_Area_FG4_(FusG):78 
    u:left_Area_FG4_(FusG):278
    u:right_Area_FG3_(FusG):79 
    u:left_Area_FG3_(FusG):279
    u:right_Dorsal_Dentate_Nucleus_(Cerebellum):80 
    u:left_Dorsal_Dentate_Nucleus_(Cerebellum):280
    u:right_Ventral_Dentate_Nucleus_(Cerebellum):81 
    u:left_Ventral_Dentate_Nucleus_(Cerebellum):281
    u:right_Area_TeI_(STG):82 
    u:left_Area_TeI_(STG):282
    u:right_Area_TI_(STG):83 
    u:left_Area_TI_(STG):283
    u:right_Area_IFS1_(IFS):84 
    u:left_Area_IFS1_(IFS):284
    u:right_Area_IFS2_(IFS):85 
    u:left_Area_IFS2_(IFS):285
    u:right_Area_IFS3_(IFS):86 
    u:left_Area_IFS3_(IFS):286
    u:right_Area_IFS4_(IFS):87 
    u:left_Area_IFS4_(IFS):287
    u:right_Area_IFJ1_(IFS,PreCS):88 
    u:left_Area_IFJ1_(IFS,PreCS):288
    u:right_Area_IFJ2_(IFS,PreCS):89 
    u:left_Area_IFJ2_(IFS,PreCS):289
    u:right_Interposed_Nucleus_(Cerebellum):90 
    u:left_Interposed_Nucleus_(Cerebellum):290
    u:right_Area_2_(PostCS):91 
    u:left_Area_2_(PostCS):291
    u:right_Ch_4_(Basal_Forebrain):92 
    u:left_Ch_4_(Basal_Forebrain):292
    u:right_Area_STS1_(STS):93 
    u:left_Area_STS1_(STS):293
    u:right_Area_STS2_(STS):94 
    u:left_Area_STS2_(STS):294
    u:right_Area_Op8_(Frontal_Operculum):95 
    u:left_Area_Op8_(Frontal_Operculum):295
    u:right_Area_Op9_(Frontal_Operculum):96 
    u:left_Area_Op9_(Frontal_Operculum):296
    u:right_Ch_123_(Basal_Forebrain):97 
    u:left_Ch_123_(Basal_Forebrain):297
    u:right_Area_6d1_(PreCG):98 
    u:left_Area_6d1_(PreCG):298
    u:right_Area_6d2_(PreCG):99 
    u:left_Area_6d2_(PreCG):299
    u:right_Area_6d3_(SFS):100
    u:left_Area_6d3_(SFS):300
    u:right_CM_(Amygdala):101
    u:left_CM_(Amygdala):301
    u:right_Area_hOc6_(POS):102
    u:left_Area_hOc6_(POS):302
    u:right_Area_hIP6_(IPS):103
    u:left_Area_hIP6_(IPS):303
    u:right_Area_hIP8_(IPS):104
    u:left_Area_hIP8_(IPS):304
    u:right_Area_hIP4_(IPS):105
    u:left_Area_hIP4_(IPS):305
    u:right_Area_hIP5_(IPS):106
    u:left_Area_hIP5_(IPS):306
    u:right_Area_hIP7_(IPS):107
    u:left_Area_hIP7_(IPS):307
    u:right_Area_hPO1_(POS):108
    u:left_Area_hPO1_(POS):308
    u:right_Area_6mp_(SMA,_mesial_SFG):109
    u:left_Area_6mp_(SMA,_mesial_SFG):309
    u:right_Area_6ma_(preSMA,_mesial_SFG):110
    u:left_Area_6ma_(preSMA,_mesial_SFG):310
    u:right_HC-Transsubiculum_(Hippocampus):111
    u:left_HC-Transsubiculum_(Hippocampus):311
    u:right_Tuberculum_(Basal_Forebrain):112
    u:left_Tuberculum_(Basal_Forebrain):312
    u:right_Terminal_islands_(Basal_Forebrain):113
    u:left_Terminal_islands_(Basal_Forebrain):313
    u:right_Area_Fo4_(OFC):114
    u:left_Area_Fo4_(OFC):314
    u:right_Area_Fo5_(OFC):115
    u:left_Area_Fo5_(OFC):315
    u:right_Area_Fo6_(OFC):116
    u:left_Area_Fo6_(OFC):316
    u:right_Area_Fo7_(OFC):117
    u:left_Area_Fo7_(OFC):317
    u:right_Area_8d1_(SFG):118
    u:left_Area_8d1_(SFG):318
    u:right_Area_8d2_(SFG):119
    u:left_Area_8d2_(SFG):319
    u:right_Area_8v2_(MFG):120
    u:left_Area_8v2_(MFG):320
    u:right_Area_8v1_(MFG):121
    u:left_Area_8v1_(MFG):321
    u:right_Area_Ig3_(Insula):122
    u:left_Area_Ig3_(Insula):322
    u:right_Area_Id4_(Insula):123
    u:left_Area_Id4_(Insula):323
    u:right_Area_Id5_(Insula):124
    u:left_Area_Id5_(Insula):324
    u:right_Area_Ia1_(Insula):125
    u:left_Area_Ia1_(Insula):325
    u:right_Area_Id6_(Insula):126
    u:left_Area_Id6_(Insula):326
    u:right_Area_SFS1_(SFS):127
    u:left_Area_SFS1_(SFS):327
    u:right_Area_SFS2_(SFS):128
    u:left_Area_SFS2_(SFS):328
    u:right_Area_MFG1_(MFG):129
    u:left_Area_MFG1_(MFG):329
    u:right_Area_MFG2_(MFG):130
    u:left_Area_MFG2_(MFG):330
    u:right_Area_Op5_(Frontal_Operculum):131
    u:left_Area_Op5_(Frontal_Operculum):331
    u:right_Area_Op6_(Frontal_Operculum):132
    u:left_Area_Op6_(Frontal_Operculum):332
    u:right_Area_Op7_(Frontal_Operculum):133
    u:left_Area_Op7_(Frontal_Operculum):333
    u:right_Area_Ph1_(PhG):134
    u:left_Area_Ph1_(PhG):334
    u:right_Area_Ph2_(PhG):135
    u:left_Area_Ph2_(PhG):335
    u:right_Area_Ph3_(PhG):136
    u:left_Area_Ph3_(PhG):336
    u:right_Area_CoS1_(CoS):137
    u:left_Area_CoS1_(CoS):337
    u:right_Medial_Accumbens,_ventral_Striatum_(Basal_Ganglia):138
    u:left_Medial_Accumbens,_ventral_Striatum_(Basal_Ganglia):338
    u:right_Fundus_of_Caudate_Nucleus,_ventral_Striatum_(Basal_Ganglia):139
    u:left_Fundus_of_Caudate_Nucleus,_ventral_Striatum_(Basal_Ganglia):339
    u:right_Fundus_of_Putamen,_Ventral_Striatum_(Basal_Ganglia):140
    u:left_Fundus_of_Putamen,_Ventral_Striatum_(Basal_Ganglia):340
    u:right_Lateral_Accumbens,_ventral_Striatum_(Basal_Ganglia):141
    u:left_Lateral_Accumbens,_ventral_Striatum_(Basal_Ganglia):341
    u:right_VP,_ventral_Pallidum_(Basal_Ganglia):142
    u:left_VP,_ventral_Pallidum_(Basal_Ganglia):342
    u:right_CGL_(Metathalamus):143
    u:left_CGL_(Metathalamus):343
    u:right_CGM_(Metathalamus):144
    u:left_CGM_(Metathalamus):344
    u:right_STN_(Subthalamus):145
    u:left_STN_(Subthalamus):345
    u:right_Area_Id9_(Insula):146
    u:left_Area_Id9_(Insula):346
    u:right_Area_Ia3_(Insula):147
    u:left_Area_Ia3_(Insula):347
    u:right_Area_Ia2_(Insula):148
    u:left_Area_Ia2_(Insula):348
    u:right_Area_Id8_(Insula):149
    u:left_Area_Id8_(Insula):349
    u:right_Area_Id10_(Insula):150
    u:left_Area_Id10_(Insula):350
    u:right_BST_(Bed_Nucleus):151
    u:left_BST_(Bed_Nucleus):351
    u:right_Frontal-I_(GapMap):152
    u:left_Frontal-I_(GapMap):352
    u:right_Frontal-II_(GapMap):153
    u:left_Frontal-II_(GapMap):353
    u:right_Temporal-to-Parietal_(GapMap):154
    u:left_Temporal-to-Parietal_(GapMap):354
    u:right_Frontal-to-Occipital_(GapMap):155
    u:left_Frontal-to-Occipital_(GapMap):355
    u:right_Frontal-to-Temporal-I_(GapMap):156
    u:left_Frontal-to-Temporal-I_(GapMap):356
    u:right_Frontal-to-Temporal-II_(GapMap):157
    u:left_Frontal-to-Temporal-II_(GapMap):357
    ----------- End regions for Julich_MNI2009c atlas --------------
    """)
            
    elif atlas == "FreeSurfer_MNI2009c":
        st.image("FS.afni.MNI2009c_asym_Atlas.png", caption="FreeSurfer_MNI2009c Atlas")
        with st.expander("View FreeSurfer_MNI2009c Region List"):
            st.text("""\
    Atlas FS.afni.MNI2009c_asym,      87 regions
    ----------- Begin regions for FS.afni.MNI2009c_asym atlas-----------
    u:Left-Cerebellum-Cortex:6  
    u:Left-Thalamus-Proper:7  
    u:Left-Caudate:8  
    u:Left-Putamen:9  
    u:Left-Pallidum:10 
    u:Brain-Stem:13 
    u:Left-Hippocampus:14 
    u:Left-Amygdala:15 
    u:Left-Accumbens-area:17 
    u:Left-VentralDC:18 
    u:Right-Cerebellum-Cortex:26 
    u:Right-Thalamus-Proper:27 
    u:Right-Caudate:28 
    u:Right-Putamen:29 
    u:Right-Pallidum:30 
    u:Right-Hippocampus:31 
    u:Right-Amygdala:32 
    u:Right-Accumbens-area:33 
    u:Right-VentralDC:34 
    u:ctx-lh-bankssts:47 
    u:ctx-lh-caudalanteriorcingulate:48 
    u:ctx-lh-caudalmiddlefrontal:49 
    u:ctx-lh-cuneus:50 
    u:ctx-lh-entorhinal:51 
    u:ctx-lh-fusiform:52 
    u:ctx-lh-inferiorparietal:53 
    u:ctx-lh-inferiortemporal:54 
    u:ctx-lh-isthmuscingulate:55 
    u:ctx-lh-lateraloccipital:56 
    u:ctx-lh-lateralorbitofrontal:57 
    u:ctx-lh-lingual:58 
    u:ctx-lh-medialorbitofrontal:59 
    u:ctx-lh-middletemporal:60 
    u:ctx-lh-parahippocampal:61 
    u:ctx-lh-paracentral:62 
    u:ctx-lh-parsopercularis:63 
    u:ctx-lh-parsorbitalis:64 
    u:ctx-lh-parstriangularis:65 
    u:ctx-lh-pericalcarine:66 
    u:ctx-lh-postcentral:67 
    u:ctx-lh-posteriorcingulate:68 
    u:ctx-lh-precentral:69 
    u:ctx-lh-precuneus:70 
    u:ctx-lh-rostralanteriorcingulate:71 
    u:ctx-lh-rostralmiddlefrontal:72 
    u:ctx-lh-superiorfrontal:73 
    u:ctx-lh-superiorparietal:74 
    u:ctx-lh-superiortemporal:75 
    u:ctx-lh-supramarginal:76 
    u:ctx-lh-frontalpole:77 
    u:ctx-lh-temporalpole:78 
    u:ctx-lh-transversetemporal:79 
    u:ctx-lh-insula:80 
    u:ctx-rh-bankssts:82 
    u:ctx-rh-caudalanteriorcingulate:83 
    u:ctx-rh-caudalmiddlefrontal:84 
    u:ctx-rh-cuneus:85 
    u:ctx-rh-entorhinal:86 
    u:ctx-rh-fusiform:87 
    u:ctx-rh-inferiorparietal:88 
    u:ctx-rh-inferiortemporal:89 
    u:ctx-rh-isthmuscingulate:90 
    u:ctx-rh-lateraloccipital:91 
    u:ctx-rh-lateralorbitofrontal:92 
    u:ctx-rh-lingual:93 
    u:ctx-rh-medialorbitofrontal:94 
    u:ctx-rh-middletemporal:95 
    u:ctx-rh-parahippocampal:96 
    u:ctx-rh-paracentral:97 
    u:ctx-rh-parsopercularis:98 
    u:ctx-rh-parsorbitalis:99 
    u:ctx-rh-parstriangularis:100
    u:ctx-rh-pericalcarine:101
    u:ctx-rh-postcentral:102
    u:ctx-rh-posteriorcingulate:103
    u:ctx-rh-precentral:104
    u:ctx-rh-precuneus:105
    u:ctx-rh-rostralanteriorcingulate:106
    u:ctx-rh-rostralmiddlefrontal:107
    u:ctx-rh-superiorfrontal:108
    u:ctx-rh-superiorparietal:109
    u:ctx-rh-superiortemporal:110
    u:ctx-rh-supramarginal:111
    u:ctx-rh-frontalpole:112
    u:ctx-rh-temporalpole:113
    u:ctx-rh-transversetemporal:114
    u:ctx-rh-insula:115
    ----------- End regions for FS.afni.MNI2009c_asym atlas --------------
    """)
    
    elif atlas == "Broadmann_MNI2009":
        st.image("Brodmann_pijn_Atlas.png", caption="Broadmann_pijn Atlas")
        with st.expander("View Broadmann_pijn Region List"):
            st.text("""\
    Atlas Brodmann_pijn,      78 regions
    ----------- Begin regions for Brodmann_pijn atlas-----------
    u:ctx-lh-BA1_3:1
    u:ctx-lh-BA2:2
    u:ctx-lh-BA4:3
    u:ctx-lh-BA5:4
    u:ctx-lh-BA6:5
    u:ctx-lh-BA7:6
    u:ctx-lh-BA8:7
    u:ctx-lh-BA9:8
    u:ctx-lh-BA10:9
    u:ctx-lh-BA11:10
    u:ctx-lh-BA13:11
    u:ctx-lh-BA16:12
    u:ctx-lh-BA17:13
    u:ctx-lh-BA18:14
    u:ctx-lh-BA19:15
    u:ctx-lh-BA20:16
    u:ctx-lh-BA21:17
    u:ctx-lh-BA22:18
    u:ctx-lh-BA23:19
    u:ctx-lh-BA24:20
    u:ctx-lh-BA25:21
    u:ctx-lh-BA26_29_30:22
    u:ctx-lh-BA27:23
    u:ctx-lh-BA28:24
    u:ctx-lh-BA31:25
    u:ctx-lh-BA32:26
    u:ctx-lh-BA33:27
    u:ctx-lh-BA34:28
    u:ctx-lh-BA35_36:29
    u:ctx-lh-BA37:30
    u:ctx-lh-BA38:31
    u:ctx-lh-BA39:32
    u:ctx-lh-BA40:33
    u:ctx-lh-BA41_42_52:34
    u:ctx-lh-BA43:35
    u:ctx-lh-BA44:36
    u:ctx-lh-BA45:37
    u:ctx-lh-BA46:38
    u:ctx-lh-BA47:39
    u:ctx-rh-BA1_3:101
    u:ctx-rh-BA2:102
    u:ctx-rh-BA4:103
    u:ctx-rh-BA5:104
    u:ctx-rh-BA6:105
    u:ctx-rh-BA7:106
    u:ctx-rh-BA8:107
    u:ctx-rh-BA9:108
    u:ctx-rh-BA10:109
    u:ctx-rh-BA11:110
    u:ctx-rh-BA13:111
    u:ctx-rh-BA16:112
    u:ctx-rh-BA17:113
    u:ctx-rh-BA18:114
    u:ctx-rh-BA19:115
    u:ctx-rh-BA20:116
    u:ctx-rh-BA21:117
    u:ctx-rh-BA22:118
    u:ctx-rh-BA23:119
    u:ctx-rh-BA24:120
    u:ctx-rh-BA25:121
    u:ctx-rh-BA26_29_30:122
    u:ctx-rh-BA27:123
    u:ctx-rh-BA28:124
    u:ctx-rh-BA31:125
    u:ctx-rh-BA32:126
    u:ctx-rh-BA33:127
    u:ctx-rh-BA34:128
    u:ctx-rh-BA35_36:129
    u:ctx-rh-BA37:130
    u:ctx-rh-BA38:131
    u:ctx-rh-BA39:132
    u:ctx-rh-BA40:133
    u:ctx-rh-BA41_42_52:134
    u:ctx-rh-BA43:135
    u:ctx-rh-BA44:136
    u:ctx-rh-BA45:137
    u:ctx-rh-BA46:138
    u:ctx-rh-BA47:139
    ----------- End regions for Brodmann_pijn atlas-----------
    """)
            
    run_button = st.sidebar.button("Run Analysis")
    if run_button:
        if cluster_file_path is not None:
            with st.spinner("Running analysis..."):
                anatomical_str, clusters = run_analysis(
                    cluster_file_path=cluster_file_path,
                    atlas=atlas,
                    task_description=task_description,
                    contrast_description=contrast_description
                )
            # Store the results in session_state so they can be accessed later
            st.session_state.anatomical_str = anatomical_str

            if st.session_state.pro_mode:
                # 1) Load your MNI template (MNI152_2009_template.nii.gz).
                template_img = nib.load("MNI152_2009_template.nii.gz")
                template_data = template_img.get_fdata()

                # 2) Get the affine. This maps voxel coordinates to MNI space.
                template_affine = template_img.affine

                            
                # 3) Compute the isosurface in voxel space.
                level = np.percentile(template_data, 90)
                verts, faces, normals, values = marching_cubes(template_data, level=level)
                # Sanitize vertices to remove any NaNs or infinite values
                verts = np.nan_to_num(verts)

                # 4) Transform isosurface vertices to MNI space using the affine.
                ones_col = np.ones((verts.shape[0], 1))
                verts_hom = np.hstack([verts, ones_col])               # shape -> (N, 4)
                verts_mni = template_affine @ verts_hom.T              # shape -> (4, N)
                verts_mni = verts_mni[:3].T                            # discard the 4th row -> (N, 3)

                # 5) Create the 3D mesh in Plotly, now in MNI space (verts_mni).
                brain_mesh = go.Mesh3d(
                    x=verts_mni[:, 0],
                    y=verts_mni[:, 1],
                    z=verts_mni[:, 2],
                    i=faces[:, 0],
                    j=faces[:, 1],
                    k=faces[:, 2],
                    opacity=0.3,
                    color='lightgrey',
                    name="Brain Surface"
                )

                # 6) Ensure cluster_coords is defined before use.
                cluster_coords = []
                for cluster in clusters:
                    # Retrieve the anatomical labels (and coordinates) for the cluster peak.
                    # Note: This call converts the cluster["Peak"] to MNI (LPI) coordinates.
                    _, _, _, mni_coords = get_anatomical_labels(cluster["Peak"], atlas, coord_system)
                    cluster_coords.append(mni_coords)

                # 7) Separate the coordinates for plotting.
                if cluster_coords:
                    cluster_x, cluster_y, cluster_z = zip(*cluster_coords)
                else:
                    cluster_x, cluster_y, cluster_z = [], [], []

                # 8) Create a Scatter3d for the clusters.
                cluster_scatter = go.Scatter3d(
                    x=cluster_x,
                    y=cluster_y,
                    z=cluster_z,
                    mode='markers',
                    marker=dict(
                        size=5,
                        color='red'
                    ),
                    name="Clusters"
                )

                fig = go.Figure(data=[brain_mesh, cluster_scatter])
                fig.update_layout(
                    title="3D Brain Visualization of your clusters (Aligned in MNI Space in LPI coordinates)",
                    scene=dict(
                        xaxis_title="X",
                        yaxis_title="Y",
                        zaxis_title="Z",
                        aspectmode="data",
                        camera=dict(
                            eye=dict(x=1.5, y=0.2, z=0.5),  # smaller values bring the camera closer
                            up=dict(x=0, y=0, z=1)
                        )
                    )
                )

                st.plotly_chart(fig, use_container_width=True)

            with st.spinner("Perplexity is thoroughly researching the internet to find the implications of these identified brain regions..."):
                if use_ai:
                    interpretation, references = synthesize_interpretation(anatomical_str, task_description, contrast_description, max_clusters)
                else:
                    interpretation, references = "AI interpretation disabled", []
            st.session_state.interpretation = interpretation
            st.session_state.references = references
            
            if use_ai:
                st.subheader(f"Deep Research Interpretation for the {max_clusters} clusters:")
            st.markdown(interpretation)
            if references:
                st.subheader("References")
                for i, ref in enumerate(references, start=1):
                    # Format: "- [1](URL): URL"
                    st.markdown(f"- {i}: {ref}")
            os.remove(cluster_file_path)
        else:
            st.warning("Please upload a cluster file (.1D) before running the analysis.")
    settings_summary = (
        f"Task Description: {task_description}\n"
        f"Analysis Description: {contrast_description}\n"
        f"Atlas: {atlas}\n"
        f"Coordinate System: {coord_system}"
    )

    # Only show the download button if the interpretation is available in session_state.
    if st.session_state.pro_mode and "interpretation" in st.session_state and st.session_state.interpretation:
        report_filename = f"{sanitize_filename(task_description)}_{sanitize_filename(atlas)}_Report.docx"
        # Generate the Word report.
        my_references = st.session_state.get("references", [])

        report_path = create_word_report(
            settings_summary,
            st.session_state.anatomical_str,
            st.session_state.interpretation,
            my_references
        )

        with open(report_path, "rb") as f:
            report_bytes = f.read()
        st.download_button(
            label="Download Report as Word Document",
            data=report_bytes,
            file_name=report_filename,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # Optionally, remove the temporary file after download.
        os.remove(report_path)
    elif "interpretation" in st.session_state and st.session_state.interpretation:
        st.warning("3d Visualization, Sonar Deep Research and Downloading the Word report are options only available only in Pro Mode. Please enable Pro Mode to access these features.")




